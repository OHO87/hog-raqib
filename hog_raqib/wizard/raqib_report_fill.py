# -*- coding: utf-8 -*-
"""تعبئة تقرير URS Client Report (docx) من بيانات التدقيق.

الملف يولده URS CMS بالدمج المسبق (اسم العميل/المواصفة/النطاق). نعبئ فقط
خلايا البيانات، ونحدد الجداول بمطابقة نص العنوان السابق لها — لا بفهرس
الجدول — لأن ترتيب الجداول يختلف بين نسخ التقرير (Stage 1 / Stage 2 /
مواقع متعددة).
"""
import base64
import io
import logging

from odoo import fields, models
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

try:
    import docx  # python-docx
except ImportError:
    docx = None
    _logger.warning("python-docx غير مثبت — تعبئة تقارير URS معطلة")


# مفاتيح تحديد الجداول: جزء مميز من نص العنوان/الخلية الأولى
KEY_MANDAYS = "Audit Man-Days Planned"
KEY_TEAM = "Lead Auditor"
KEY_SITE = "Full Address of the Permanent Site"
KEY_PROCESSES = "Name of Process or Department"
KEY_REGULATORY = "Regulatory Awareness"
KEY_STAGE2_READY = "Reviewed for Stage 2 Readiness"
KEY_COMMENTS = "Type of Comment"
KEY_NEXT_VISIT = "Planned Month/Year"


class RaqibReportFill(models.TransientModel):
    _name = "raqib.report.fill"
    _description = "تعبئة تقرير URS"

    audit_id = fields.Many2one("raqib.audit", required=True)
    standard_id = fields.Many2one(
        "raqib.standard", string="المواصفة",
        help="في التدقيق متعدد المواصفات: يولد تقرير هذه المواصفة فقط "
             "(بنودها وملاحظاتها). اتركه فارغاً لتقرير يشمل الكل.")
    template_file = fields.Binary("ملف التقرير من CMS (docx)", required=True)
    template_filename = fields.Char("اسم الملف")

    def _report_lines(self):
        """بنود الفحص الداخلة في هذا التقرير (مفلترة بالمواصفة إن حددت)."""
        audit = self.audit_id
        lines = audit.line_ids.filtered(lambda l: l.result != "pending")
        if self.standard_id:
            std = self.standard_id
            lines = lines.filtered(
                lambda l: std in l._covered_clauses().mapped("standard_id"))
        return lines

    def _report_findings(self):
        audit = self.audit_id
        findings = audit.finding_ids
        if self.standard_id:
            std = self.standard_id
            findings = findings.filtered(
                lambda f: std in (f.standard_ids
                                  or f.clause_id.standard_id))
        return findings

    # ------------------------------------------------------------------
    # أدوات مساعدة على بنية الوثيقة
    # ------------------------------------------------------------------
    @staticmethod
    def _table_text(table, max_cells=30):
        txt = []
        for row in table.rows[:3]:
            for cell in row.cells[:max_cells]:
                txt.append(cell.text)
        return " | ".join(txt)

    def _find_table(self, document, key):
        for table in document.tables:
            if key.lower() in self._table_text(table).lower():
                return table
        return None

    @staticmethod
    def _set_cell(cell, text):
        """كتابة نص في خلية مع الحفاظ على تنسيق أول Run إن وجد."""
        if text is None:
            return
        text = str(text)
        # امسح الفقرات الزائدة واكتب في الأولى
        first = cell.paragraphs[0]
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        if first.runs:
            first.runs[0].text = text
            for r in first.runs[1:]:
                r.text = ""
        else:
            first.add_run(text)

    @staticmethod
    def _row_is_empty(row, skip_first=1):
        return all(not c.text.strip() for c in row.cells[skip_first:])

    @staticmethod
    def _add_row_like_last(table):
        import copy
        new_tr = copy.deepcopy(table.rows[-1]._tr)
        for tc in new_tr.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
            tc.text = ""
        table.rows[-1]._tr.addnext(new_tr)
        return table.rows[-1]

    def _fill_repeating(self, table, rows_data, header_rows=1):
        """تعبئة صفوف متكررة: تُستخدم الصفوف الفارغة الجاهزة أولاً ثم تضاف صفوف."""
        data_rows = [r for r in table.rows[header_rows:]]
        idx = 0
        for values in rows_data:
            row = None
            while idx < len(data_rows):
                cand = data_rows[idx]
                idx += 1
                if self._row_is_empty(cand, skip_first=0) or self._row_is_empty(cand):
                    row = cand
                    break
            if row is None:
                row = self._add_row_like_last(table)
            for ci, val in enumerate(values):
                if ci < len(row.cells):
                    self._set_cell(row.cells[ci], val)

    # ------------------------------------------------------------------
    # التعبئة
    # ------------------------------------------------------------------
    def action_fill(self):
        self.ensure_one()
        if docx is None:
            raise UserError(
                "python-docx غير مثبت على الخادم. ثبّته بـ: pip3 install python-docx "
                "ثم أعد تشغيل أودو.")
        audit = self.audit_id
        document = docx.Document(io.BytesIO(base64.b64decode(self.template_file)))
        filled, missing = [], []

        # 1) Man-days
        t = self._find_table(document, KEY_MANDAYS)
        if t is not None:
            labels = {
                "man-days planned": audit.mandays_planned,
                "actual audit man-days": audit.mandays_actual,
                "difference": audit.mandays_diff,
                "justify": audit.mandays_justification or "",
                "end date": audit.date_end and audit.date_end.strftime("%d/%m/%Y") or "",
            }
            for row in t.rows:
                head = row.cells[0].text.lower()
                for key, val in labels.items():
                    if key in head and len(row.cells) > 1:
                        self._set_cell(row.cells[-1], val)
            filled.append("Man-Days")
        else:
            missing.append("Man-Days")

        # 2) فريق التدقيق
        t = self._find_table(document, KEY_TEAM)
        if t is not None:
            members = list(audit.team_ids)
            lead = [m for m in members if m.role == "lead"]
            others = [m for m in members if m.role == "member"]
            for row in t.rows:
                head = row.cells[0].text.lower()
                if "lead auditor" in head and lead:
                    self._set_cell(row.cells[1], lead[0].name)
                    if len(row.cells) > 2:
                        self._set_cell(row.cells[2], lead[0].days)
                for i in range(1, 5):
                    if ("member %d" % i) in head and len(others) >= i:
                        self._set_cell(row.cells[1], others[i - 1].name)
                        if len(row.cells) > 2:
                            self._set_cell(row.cells[2], others[i - 1].days)
                if "specialist" in head:
                    sp = [m for m in members if m.role == "specialist"]
                    if sp:
                        self._set_cell(row.cells[-1], sp[0].name)
                if "translator" in head:
                    tr = [m for m in members if m.role == "translator"]
                    if tr:
                        self._set_cell(row.cells[-1], tr[0].name)
            filled.append("Audit Team")
        else:
            missing.append("Audit Team")

        # 3) الموقع
        t = self._find_table(document, KEY_SITE)
        if t is not None:
            client = audit.client_id
            for row in t.rows:
                head = row.cells[0].text.lower()
                if "full address" in head:
                    self._set_cell(row.cells[-1], client.address or "")
                elif "total employees" in head:
                    self._set_cell(row.cells[-1],
                                   audit.employees_onsite or client.employee_count or "")
                elif "full-time" in head or "equivalent" in head:
                    self._set_cell(row.cells[-1],
                                   audit.fte_equiv or client.fte_equiv or "")
            filled.append("Site")
        else:
            missing.append("Site")

        # 4) العمليات الجوهرية
        t = self._find_table(document, KEY_PROCESSES)
        if t is not None:
            rows = [(p.reference or str(i + 1), p.name)
                    for i, p in enumerate(audit.client_id.process_ids)]
            if rows:
                self._fill_repeating(t, rows)
            filled.append("Fundamental Processes")
        else:
            missing.append("Fundamental Processes")

        # 5) الوعي التنظيمي
        t = self._find_table(document, KEY_REGULATORY)
        if t is not None:
            client = audit.client_id
            for row in t.rows:
                head = row.cells[0].text.lower()
                if "regulatory awareness" in head and len(row.cells) > 1:
                    self._set_cell(row.cells[-1], client.regulatory_text or "")
                elif "comment" in head and len(row.cells) > 1:
                    self._set_cell(row.cells[-1], client.regulatory_comment or "")
            filled.append("Regulatory")
        else:
            missing.append("Regulatory")

        # 6) جاهزية Stage 2: سطر لكل بند مفحوص
        t = self._find_table(document, KEY_STAGE2_READY)
        if t is not None:
            result_label = dict(
                audit.line_ids._fields["result"]._description_selection(self.env))
            rows = []
            for line in self._report_lines():
                ref = line.doc_reference or ""
                if line.last_review_date:
                    ref = ("%s — آخر مراجعة %s" % (
                        ref, line.last_review_date.strftime("%d/%m/%Y"))).strip(" —")
                rows.append((
                    "%s %s" % (line.number, line.clause_name),
                    ref + ((" | " + line.note) if line.note else ""),
                    result_label.get(line.result, line.result),
                ))
            if rows:
                self._fill_repeating(t, rows)
            filled.append("Stage 2 Readiness")
        else:
            missing.append("Stage 2 Readiness")

        # 7) الملاحظات Comments Raised
        t = self._find_table(document, KEY_COMMENTS)
        if t is not None:
            rows = [(str(f.number), f.description, f.urs_type_label())
                    for f in self._report_findings()]
            if rows:
                self._fill_repeating(t, rows)
            filled.append("Comments Raised")
        else:
            missing.append("Comments Raised")

        # 8) خطة الزيارة القادمة (الصف الرئيسي فقط)
        t = self._find_table(document, KEY_NEXT_VISIT)
        if t is not None:
            for row in t.rows:
                head = row.cells[0].text.lower()
                if "planned month" in head and len(row.cells) > 1:
                    self._set_cell(row.cells[-1], audit.next_visit_month or "")
                if "activity type" in head and len(row.cells) > 1:
                    self._set_cell(row.cells[-1], audit.next_visit_activity or "")
            filled.append("Next Visit")
        else:
            missing.append("Next Visit")

        out = io.BytesIO()
        document.save(out)
        data = base64.b64encode(out.getvalue())
        suffix = (" - %s" % self.standard_id.code) if self.standard_id else ""
        fname = (self.template_filename or "urs_report.docx").replace(
            ".docx", "") + suffix + " - FILLED.docx"
        audit.write({
            "report_docx": data,
            "report_docx_name": fname,
        })
        # أرشفة نسخة دائمة لكل مواصفة (التدقيق متعدد المواصفات يولد عدة تقارير)
        self.env["ir.attachment"].create({
            "name": fname,
            "res_model": "raqib.audit",
            "res_id": audit.id,
            "datas": data,
            "mimetype": ("application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document"),
        })
        audit.message_post(body=(
            "تم توليد التقرير%s. جداول عبئت: %s.%s"
            % ((" (مواصفة %s)" % self.standard_id.code)
               if self.standard_id else "",
               ", ".join(filled),
               (" لم يعثر على: %s — عبئها يدوياً." % ", ".join(missing))
               if missing else "")))
        return {
            "type": "ir.actions.act_url",
            "url": "/web/content/raqib.audit/%d/report_docx/%s?download=true"
                   % (audit.id, fname),
            "target": "self",
        }
